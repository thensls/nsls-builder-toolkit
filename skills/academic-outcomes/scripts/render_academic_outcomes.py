#!/usr/bin/env python3.12
"""
Render an academic learner-outcomes document (.docx) from structured JSON.

Partner-neutral: this is the WGU-crosswalk renderer (build_society_wgu.py,
2026-06-23) with the WGU-specific crosswalk column removed. It renders both a
single-track doc and a bundle (many tracks) from ONE document shape, so per-track
docs and the bundle never drift.

Usage:
    python3.12 render_academic_outcomes.py <document.json> <out.docx>

Then upload the .docx as a Google Doc with gdoc-build's `gws drive files create
--upload ...` step (Society branding is already baked into this renderer).

Requires python-docx. Per the NSLS env note:
    [ -d /tmp/pptx_deps/docx ] || python3.12 -m pip install python-docx --target /tmp/pptx_deps -q

Document JSON shape — see references/outcomes-json-schema.md. In brief:
{
  "title": "...", "subtitle": "...",         # required
  "about": "...",                            # optional paragraph
  "how_learning_works": ["**Lead.** rest", ...],  # optional bullets
  "overlap": {"headers": [...], "rows": [[...]]},  # optional summary table
  "sections": [                              # required; 1+ sections
    {"heading": "Learning Experiences", "items": [ <ITEM>, ... ]}
  ]
}
Each ITEM is the per-track structured record stored in Airtable
`academic_outcomes_json`:
{
  "num": 1, "title": "Career Clarity",
  "meta": "Track · Live · ~3.5 hours · Develops Self-Awareness.",
  "description": "...",
  "comp_label": "Competencies",
  "comp_headers": ["Competency", "Evidence the member produces"],
  "comp_rows": [["...", "..."], ...],
  "how_it_works": "...",
  "excluded": "..." | null,
  "outputs": "...",
  "status": "Status: live."
}
"""
import sys
import json

sys.path.insert(0, "/tmp/pptx_deps")

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Society brand palette (from build_society_wgu.py).
P = {
    "header_bg": "3D2F1F",
    "header_fg": (0xF5, 0xE6, 0xCB),
    "row_band": "FAF4E8",
    "h1": (0x3D, 0x2F, 0x1F),
    "h2": (0x6B, 0x4F, 0x2F),
}


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(*P["header_fg"])
        run.font.size = Pt(10.5)
        set_cell_bg(hdr[i], P["header_bg"])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        bg = P["row_band"] if r_idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row_data):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    color = P["h1"] if level <= 1 else P["h2"]
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_para(doc, text, italic=False):
    p = doc.add_paragraph()
    for i, chunk in enumerate(text.split("**")):
        run = p.add_run(chunk)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
        if italic:
            run.italic = True
            run.font.size = Pt(10)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    for i, chunk in enumerate(text.split("**")):
        run = p.add_run(chunk)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


def render_item(doc, item):
    """Render one track / capability. Partner-neutral: NO crosswalk table."""
    num = item.get("num")
    heading = f"{num}. {item['title']}" if num else item["title"]
    add_heading(doc, heading, 2)
    if item.get("meta"):
        add_para(doc, item["meta"], italic=True)
    add_para(doc, f"**Description.** {item['description']}")
    comp_label = item.get("comp_label", "Competencies")
    add_para(doc, f"**{comp_label}**")
    headers = item.get("comp_headers", ["Competency", "Evidence the member produces"])
    # Two-column competency tables get the wide split; leveled (3-col) tables get
    # a narrower first column.
    if len(headers) == 2:
        widths = [Inches(3.4), Inches(3.3)]
    elif len(headers) == 3:
        widths = [Inches(1.1), Inches(3.0), Inches(2.6)]
    else:
        widths = None
    add_table(doc, headers, item["comp_rows"], widths)
    add_para(doc, f"**How it works.** {item['how_it_works']}")
    if item.get("excluded"):
        add_para(doc, f"**Excluded by design.** {item['excluded']}")
    add_para(doc, f"**Outputs.** {item['outputs']}")
    if item.get("status"):
        add_para(doc, item["status"], italic=True)
    doc.add_paragraph()


def build(document):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- TITLE ----
    title = doc.add_paragraph()
    tr = title.add_run(document["title"])
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(*P["h1"])
    if document.get("subtitle"):
        sub = doc.add_paragraph()
        sr = sub.add_run(document["subtitle"])
        sr.font.size = Pt(12)
        sr.italic = True
        sr.font.color.rgb = RGBColor(*P["h2"])

    # ---- ABOUT ----
    if document.get("about"):
        add_heading(doc, "About this package", 1)
        add_para(doc, document["about"])

    # ---- HOW LEARNING WORKS ----
    if document.get("how_learning_works"):
        add_heading(doc, "How learning works on Society", 1)
        add_para(doc, "Every Society experience runs the same model:")
        for b in document["how_learning_works"]:
            add_bullet(doc, b)

    # ---- OVERLAP AT A GLANCE ----
    overlap = document.get("overlap")
    if overlap and overlap.get("rows"):
        add_heading(doc, "At a glance", 1)
        add_table(doc, overlap["headers"], overlap["rows"])

    # ---- SECTIONS ----
    for sec in document.get("sections", []):
        if sec.get("heading"):
            add_heading(doc, sec["heading"], 1)
        for item in sec.get("items", []):
            render_item(doc, item)

    return doc


def main():
    if len(sys.argv) != 3:
        sys.exit("usage: render_academic_outcomes.py <document.json> <out.docx>")
    with open(sys.argv[1]) as f:
        document = json.load(f)
    doc = build(document)
    doc.save(sys.argv[2])
    print(f"Saved {sys.argv[2]}")


if __name__ == "__main__":
    main()
