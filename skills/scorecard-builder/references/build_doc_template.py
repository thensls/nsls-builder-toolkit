"""Render a founder-template ScoreCard as an NSLS-branded Google Doc.

USAGE
  1. cp this file to ~/build_<name>_scorecard.py
  2. Fill the BODY section (marked ==== FILL THIS ====). Keep every uncertain
     value as a [bracket] — those are confirmed WITH the report, not by you.
  3. PYTHONPATH="$HOME/.local/lib/nsls-pydeps:/tmp/pptx_deps" python3.12 ~/build_<name>_scorecard.py
     (if import fails: python3.12 -m pip install --upgrade --force-reinstall \
      python-docx --target ~/.local/lib/nsls-pydeps -q)
     The script derives the output filename from NAME (e.g. ~/chelsea_byers_scorecard.docx)
     and PRINTS both commands you need — upload, then the REQUIRED HR share.
  4. Run the printed STEP 1 upload command (it fills in the right filename):
       cd ~ && gws drive files create \
         --json '{"name":"ScoreCard — <Name> — <Role> — <FY> (DRAFT)","mimeType":"application/vnd.google-apps.document"}' \
         --upload <the-derived-filename>.docx \
         --upload-content-type "application/vnd.openxmlformats-officedocument.wordprocessingml.document" \
         --format json | grep -v -i keyring | tail -6
  5. Run the printed STEP 2 share command — REQUIRED. Every card goes to HR
     (Jenna Fontanez, jfontanez@nsls.org) as commenter at render time. An
     unshared card is an orphan: it never reaches the system of record.
     Note the param split — role/type/emailAddress are the BODY (--json);
     sendNotificationEmail/emailMessage are QUERY params (--params).
  6. Return the URL to the manager, state that the HR share is done, and name
     what's still theirs: share with the report, resolve the [brackets], then
     retitle (DRAFT) -> (FINAL) and notify HR. This skill NEVER writes Airtable.

Helpers below produce real Word elements that survive .docx -> Google Doc
conversion. add_runs() parses mini-markdown inside any paragraph or table cell:
**bold**, *italic*, `code`, and \\n line breaks.
"""
import os, sys, re
sys.path.insert(0, '/tmp/pptx_deps')  # legacy location — /tmp cleanup can gut it
sys.path.insert(0, os.path.expanduser('~/.local/lib/nsls-pydeps'))  # durable home, wins
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- BRAND (NSLS navy) ----
P = {"header_bg": "1A2B4A", "header_fg": (0xFF, 0xFF, 0xFF),
     "row_band": "F5F7FA", "h1": (0x1A, 0x2B, 0x4A), "h2": (0x2B, 0x4A, 0x7A)}

# ---- HELPERS ----
def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color='BFBFBF', size='6'):
    # Explicit per-cell borders: Google's .docx importer DROPS the borders the
    # built-in table.style promises, so style-only tables land borderless.
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), size); b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)

TOKEN = re.compile(r'(\*\*.+?\*\*|`.+?`|\*.+?\*)', re.S)

def add_runs(p, text, size=11):
    """Render **bold**, *italic*, `code`, and \\n line breaks into runs."""
    for tok in TOKEN.split(text):
        if not tok:
            continue
        bold = italic = code = False
        if tok.startswith('**') and tok.endswith('**'):
            tok = tok[2:-2]; bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            tok = tok[1:-1]; code = True
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 1:
            tok = tok[1:-1]; italic = True
        for i, seg in enumerate(tok.split('\n')):
            if i > 0:
                p.add_run().add_break()
            r = p.add_run(seg); r.font.size = Pt(size)
            r.bold = bold; r.italic = italic
            if code:
                r.font.name = 'Consolas'; r.font.size = Pt(9.5)

def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*(P["h1"] if level <= 1 else P["h2"]))
    return h

def para(text, italic=False, size=11):
    p = doc.add_paragraph(); add_runs(p, text, size=size)
    if italic:
        for r in p.runs:
            r.italic = True
    return p

def bullet(text):
    add_runs(doc.add_paragraph(style='List Bullet'), text)

def numbered(text):
    add_runs(doc.add_paragraph(style='List Number'), text)

def rich_table(headers, rows, widths=None, mono_first=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(*P["header_fg"]); r.font.size = Pt(10)
        set_cell_bg(c, P["header_bg"]); set_cell_borders(c); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, rowdata in enumerate(rows):
        cells = t.add_row().cells
        bg = P["row_band"] if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(rowdata):
            cells[i].text = ''
            add_runs(cells[i].paragraphs[0], val, size=10)
            if mono_first and i == 0:
                for r in cells[i].paragraphs[0].runs:
                    r.font.name = 'Consolas'
            set_cell_bg(cells[i], bg); set_cell_borders(cells[i]); cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if widths:
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Cm(w)
    return t

def rule():
    pPr = doc.add_paragraph()._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), 'BFBFBF')
    pbdr.append(b); pPr.append(pbdr)

# ---- DOC SETUP ----
doc = Document()
n = doc.styles['Normal']; n.font.name = 'Calibri'; n.font.size = Pt(11)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8)
    s.left_margin = s.right_margin = Cm(1.9)

# ============================ FILL THIS ============================
# Replace every value below. Keep [brackets] for anything the report confirms.

NAME  = "[Employee Name]"
ROLE  = "[Role Title]"
FY    = "FY2026"

# --- Title ---
t = doc.add_paragraph().add_run(f"ScoreCard — {NAME}")
t.bold = True; t.font.size = Pt(20); t.font.color.rgb = RGBColor(*P["h1"])
sr = doc.add_paragraph().add_run(f"{ROLE} · {FY}")
sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(*P["h2"])
# Assert only what stays true for the life of the Doc. DRAFT/FINAL status lives
# in the TITLE (retitled on confirmed alignment), so don't restate it here — and
# never claim a share state: the skill shares this with HR the moment it renders.
tg = doc.add_paragraph().add_run(
    "Alignment instrument for manager↔employee — current status is in the document title. "
    "Shared with HR for tracking; never written to Airtable by this skill."
)
tg.italic = True; tg.font.size = Pt(10.5); tg.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
para("**[bracketed]** = draft placeholder — the employee confirms or changes each in the alignment conversation.", size=10)
rule()

# --- Mission ---
heading("Mission", 1)
para("**This role exists to [one-sentence reason the seat is in the org].**")

# --- Side A: Accountabilities (3-6, weights sum to 100) ---
heading(f"Side A — Accountabilities (weighted to 100%)", 1)
para("Each accountability is a durable **outcome**; the **Pass/Fail Measure** is the binary \"Meets\" bar. A = clearly beat it · B = met it · C = missed.")
rich_table(
    ["#", "Accountability (outcome)", "Wt", "Pass/Fail Measure (binary \"Meets\" bar)"],
    [
        ["1", "**[Outcome — not a task]** — [what great looks like]", "[30%]", "[Concrete binary bar with a target + date, e.g. X reaches [target] by [date]]"],
        ["2", "**[Outcome]**", "[20%]", "[Binary bar]"],
        ["3", "**[Outcome]**", "[15%]", "[Binary bar]"],
        ["4", "**[Outcome]**", "[20%]", "[Binary bar]"],
        ["5", "**[Outcome]**", "[15%]", "[Binary bar]"],
        ["", "**Total**", "**100%**", ""],
    ],
    widths=[0.9, 7.2, 1.1, 7.6], mono_first=True)
para("*Rating: A = predictably meets/exceeds · B = predictably somewhat short · C = predictably far short.*", size=10)
# (Optional) portfolio bullets for a coordination role — see portfolio-2x2-framework.md
# para("**Draft portfolio for Accountability #1** *(confirm together):*", size=10.5)
# bullet("**Growth driver** — [project] [date]")
# bullet("**Operating efficiency** — [project] [date]")
# bullet("**Hygiene** — [project] [date]")
# bullet("**Reliability** — [project] [date]")
rule()

# --- Core Values: BINARY Pass/Flag gate (no ±15% modifier) ---
heading("Core Values", 1)
para("*(binary Pass / Flag gate against each MAR — not part of the 100%, not a modifier)*", italic=True, size=10)
rich_table(["Core Value", "Definition", "MAR", "Pass / Flag"],
    [["Mission Driven", "Help the greatest number of people in the most meaningful way.", "4", "—"],
     ["Accountability", "Hold ourselves and others responsible for integrity, commitments, and results through courageous honesty.", "4", "—"],
     ["Get it Done, Make it Fun", "Deliver results while fostering a positive, fun, and collaborative work environment.", "4", "—"],
     ["Lead with Heart & Service", "Lead with empathy, authenticity, and humility through heart-centered, servant leadership.", "4", "—"],
     ["Continuous Improvement & Innovation", "Prioritize progress over perfection to improve yourself, our offerings, and the company.", "4", "—"]],
    widths=[3.6, 9.4, 1.2, 1.8])
rule()

# --- Side B: 5-8 competencies from the Competency Bank ---
heading("Side B — Core Competencies", 1)
para("*(5–8 from the Competency Bank; MAR set by role)*", italic=True, size=10)
rich_table(["Competency", "Group", "Definition", "MAR", "Rating"],
    [["[Competency]", "[Group]", "[Definition from the bank]", "[4]", "—"],
     ["[Competency]", "[Group]", "[Definition]", "[4]", "—"],
     ["[Competency]", "[Group]", "[Definition]", "[3]", "—"]],
    widths=[3.4, 2.2, 7.2, 1.2, 1.8])
para("*Rating: 4 = Excellent · 3 = Very Good · 2 = Good · 1 = Weak.*", size=10)
para("**Growth Focus (proposed): [Competency]** — [one specific observable action to practice this period].")
rule()

# --- Open questions + alignment notes (talk-through material) ---
heading("Open questions — to talk through together", 1)
bullet("**[Question the manager and report should decide together, not pre-decided here].**")

heading("Alignment notes — for the conversation", 1)
numbered("**[Context the report should know — a watch-item, a reassignment, a delegation note].**")
numbered("**Bracketed values are drafts.** Every [ … ] is confirmed in this conversation.")
# ========================== END FILL THIS ==========================

import os
# Derive the output filename from NAME so it matches the upload command exactly.
slug = re.sub(r'[^a-z0-9]+', '_', NAME.lower()).strip('_') or "scorecard"
fname = f"{slug}_scorecard.docx"
out = os.path.expanduser(f"~/{fname}")
doc.save(out)
print("saved", out)
print("\nSTEP 1 — upload it with (run from ~):")
# `set -o pipefail` on every printed command: these all end in `| grep | tail`,
# so without it the pipeline's exit status is tail's (0) and a gws 4xx reads as
# success. That turns a failed share into a silently orphaned ScoreCard.
print(
    f'set -o pipefail; cd ~ && gws drive files create '
    f'--json \'{{"name":"ScoreCard — {NAME} — {ROLE} — {FY} (DRAFT)",'
    f'"mimeType":"application/vnd.google-apps.document"}}\' '
    f'--upload {fname} '
    f'--upload-content-type "application/vnd.openxmlformats-officedocument.wordprocessingml.document" '
    f'--format json | grep -v -i keyring | tail -6'
)

# STEP 2 is not optional. An unshared ScoreCard never reaches the system of
# record — it dies in the manager's Drive and the work is redone next cycle.
# Jenna Fontanez owns the scorecard process; commenter, not writer, because the
# manager and report own the card's content.
print("\nSTEP 2 — REQUIRED: share it with HR (Jenna). Fill in <DOC_ID> from step 1,")
print("and replace <NOTE> with: whose card, who the manager is, DRAFT vs FINAL,")
print("plus (revise mode) the change summary + triage verdict + prior card URL.")
print(
    f'set -o pipefail; cd ~ && gws drive permissions create '
    f'--params \'{{"fileId":"<DOC_ID>","sendNotificationEmail":true,'
    f'"emailMessage":"<NOTE>"}}\' '
    f'--json \'{{"role":"commenter","type":"user",'
    f'"emailAddress":"jfontanez@nsls.org"}}\' '
    f'| grep -v -i keyring | tail -5'
)
print("\nSTEP 3 — confirm it landed (exit code alone is not proof):")
print('gws drive permissions list --params \'{"fileId":"<DOC_ID>"}\''
      ' | grep -v -i keyring | grep -i fontanez')
print("\nIf STEP 2 or 3 fails, do NOT report the card as done — give the manager")
print("the manual fallback (Share -> jfontanez@nsls.org -> Commenter). See SKILL.md.")
