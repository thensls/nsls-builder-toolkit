"""Starter script for building a nicely-formatted NSLS / Society Google Doc.

Usage:
    1. cp this file to ~/build_<short-name>.py and customize the BODY section.
    2. PYTHONPATH=/tmp/pptx_deps python3.12 ~/build_<short-name>.py
    3. cd ~ && gws drive files create \\
         --json '{"name":"<doc title>","mimeType":"application/vnd.google-apps.document"}' \\
         --upload <short-name>.docx \\
         --upload-content-type "application/vnd.openxmlformats-officedocument.wordprocessingml.document" \\
         --format json | tail -10
    4. Open the returned URL.
    5. rm ~/build_<short-name>.py ~/<short-name>.docx

Helpers below produce real Word elements that survive .docx -> Google Doc
conversion: styled tables, heading hierarchy, code blocks, bullet/numbered
lists, cell shading, alternating row banding.

DO NOT switch to pandoc-from-markdown. See SKILL.md "Common Mistakes."
"""

import sys
sys.path.insert(0, '/tmp/pptx_deps')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---- BRAND CONSTANTS -------------------------------------------------
# Switch BRAND between "NSLS" and "Society" for different palettes.
BRAND = "NSLS"

PALETTES = {
    "NSLS": {
        "header_bg": "1A2B4A",      # navy
        "header_fg": (0xFF, 0xFF, 0xFF),
        "row_band": "F5F7FA",
        "h1": (0x1A, 0x2B, 0x4A),
        "h2": (0x2B, 0x4A, 0x7A),
        "code_bg": "F5F7FA",
    },
    "Society": {
        "header_bg": "3D2F1F",      # deep brown
        "header_fg": (0xF5, 0xE6, 0xCB),  # cream
        "row_band": "FAF4E8",
        "h1": (0x3D, 0x2F, 0x1F),
        "h2": (0x6B, 0x4F, 0x2F),
        "code_bg": "FAF4E8",
    },
}
P = PALETTES[BRAND]


# ---- HELPERS ---------------------------------------------------------

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color='BFBFBF', size='6'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_table(doc, headers, rows, col_widths=None):
    """Real Word table with NSLS-styled header row + alternating banding."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'  # built-in, gives borders on import
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''   # IMPORTANT: clear default content before adding run
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(*P["header_fg"])
        run.font.size = Pt(10.5)
        set_cell_bg(hdr[i], P["header_bg"])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # body rows with alternating banding
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        bg = P["row_band"] if r_idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                run.font.name = 'Consolas'
            set_cell_bg(cells[i], bg)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_heading(doc, text, level):
    """Use built-in Heading 1/2/3 — Google Docs maps these cleanly."""
    h = doc.add_heading(text, level=level)
    color = P["h1"] if level <= 1 else P["h2"]
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def add_para(doc, text, italic=False):
    """Paragraph with **bold** segments. Splits on ** for inline emphasis."""
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, chunk in enumerate(parts):
        run = p.add_run(chunk)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
        if italic:
            run.italic = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')   # exact built-in name
    parts = text.split('**')
    for i, chunk in enumerate(parts):
        run = p.add_run(chunk)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_code(doc, text):
    """Code block: monospace + light brand-grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), P["code_bg"])
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*P["h1"])
    return p


# ---- DOC SETUP -------------------------------------------------------

doc = Document()

# Body font default
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


# ---- BODY (CUSTOMIZE THIS SECTION) -----------------------------------

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title.add_run('Your Doc Title Goes Here')
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = RGBColor(*P["h1"])

add_heading(doc, 'First Section', 1)
add_para(doc, "Body paragraph. Use **double asterisks** to bold inline.")

add_heading(doc, 'A Subsection', 2)
add_para(doc, "Lead with the answer. Plain language. Short sentences.")
add_bullet(doc, "Bullet one — keep these scannable.")
add_bullet(doc, "Bullet two with **bold inline** for emphasis.")

add_heading(doc, 'A Table', 2)
add_table(
    doc,
    headers=['Column A', 'Column B'],
    rows=[
        ['First cell', 'Second cell'],
        ['Third cell', 'Fourth cell'],
    ],
    col_widths=[Inches(2.0), Inches(4.5)],
)

add_heading(doc, 'A Numbered List', 2)
add_numbered(doc, "Step one")
add_numbered(doc, "Step two")
add_numbered(doc, "Step three")

add_heading(doc, 'A Code Block', 2)
add_code(doc, "git checkout main && git pull origin main")


# ---- SAVE ------------------------------------------------------------

output_path = '/Users/k/your_doc_name.docx'   # MUST be in ~ for gws --upload
doc.save(output_path)
print(f"Saved {output_path}")
print("Next: cd ~ && gws drive files create --json '{\"name\":\"...\",\"mimeType\":\"application/vnd.google-apps.document\"}' --upload your_doc_name.docx --upload-content-type 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' --format json | tail -10")
